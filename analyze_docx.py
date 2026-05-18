from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document


def configure_stdio() -> None:
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if stream is not None and hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def business_table_label(index_zero_based: int) -> str:
    return f"Bảng {index_zero_based + 1}"


def business_row_label(index_zero_based: int) -> str:
    return f"Hàng {index_zero_based + 1}"


def business_col_label(index_zero_based: int) -> str:
    return f"Cột {index_zero_based + 1}"


def cell_merge_info(cell) -> tuple[int, str | None]:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1, None

    grid_span = getattr(tc_pr, "gridSpan", None)
    v_merge = getattr(tc_pr, "vMerge", None)

    span = 1
    if grid_span is not None and getattr(grid_span, "val", None):
        try:
            span = int(grid_span.val)
        except (TypeError, ValueError):
            span = 1

    vmerge = None
    if v_merge is not None:
        vmerge = getattr(v_merge, "val", None) or "continue"

    return span, vmerge


def dump_paragraphs(doc: Document) -> None:
    print("== Paragraphs ==")
    for idx, para in enumerate(doc.paragraphs):
        text = normalize_text(para.text)
        print(f"P{idx:02d} | Paragraph {idx + 1} | len={len(text):03d} | {text}")


def dump_tables(doc: Document) -> None:
    print("== Tables ==")
    for t_idx, table in enumerate(doc.tables):
        print(f"T{t_idx:02d} | {business_table_label(t_idx)} | rows={len(table.rows)}")
        for r_idx, row in enumerate(table.rows):
            parts: list[str] = []
            for c_idx, cell in enumerate(row.cells):
                text = normalize_text(cell.text)
                span, vmerge = cell_merge_info(cell)
                extra = []
                if span > 1:
                    extra.append(f"span={span}")
                if vmerge:
                    extra.append(f"vMerge={vmerge}")
                meta = f" ({', '.join(extra)})" if extra else ""
                business = f"{business_row_label(r_idx)}, {business_col_label(c_idx)}"
                parts.append(f"R{r_idx}C{c_idx} [{business}]{meta}={text}")
            print("  " + " | ".join(parts))


def dump_body_order(doc: Document) -> None:
    print("== Body Order ==")
    p_idx = 0
    t_idx = 0

    for node in doc.element.body.iterchildren():
        if node.tag.endswith("}p"):
            para = doc.paragraphs[p_idx]
            text = normalize_text(para.text)
            print(f"BODY P{p_idx:02d} | Paragraph {p_idx + 1} | {text}")
            p_idx += 1
        elif node.tag.endswith("}tbl"):
            print(f"BODY T{t_idx:02d} | {business_table_label(t_idx)}")
            t_idx += 1


def analyze_file(path: Path) -> None:
    print(f"\n===== {path.name} =====")
    doc = Document(path)
    dump_body_order(doc)
    dump_paragraphs(doc)
    dump_tables(doc)


def main() -> None:
    configure_stdio()
    parser = argparse.ArgumentParser(
        description="In cấu trúc paragraphs/tables của DOCX để đối chiếu mapping."
    )
    parser.add_argument("files", nargs="+", help="Đường dẫn một hoặc nhiều file .docx")
    args = parser.parse_args()

    for raw_path in args.files:
        analyze_file(Path(raw_path))


if __name__ == "__main__":
    main()
