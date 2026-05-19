from io import BytesIO
import os

import pytest
from docx import Document

os.environ.setdefault("DATABASE_URL", "sqlite:///./test.db")

from app.models.vessel import InspectionTypeEnum, LengthGroupEnum, MaterialEnum
from app.services.docx_parser import (
    DocxParserError,
    extract_province_code,
    get_length_group,
    parse_docx,
)


@pytest.fixture
def docx_factory(tmp_path):
    """Tạo file DOCX tạm từ nội dung text để test parser như luồng upload thật."""

    def _create(name: str, text_content: str):
        document = Document()
        for line in text_content.strip().splitlines():
            document.add_paragraph(line)

        path = tmp_path / name
        document.save(path)
        return path

    return _create


@pytest.fixture
def table_docx_factory(tmp_path):
    """Tạo file DOCX dạng bảng key/value để kiểm tra parser đọc được table cells."""

    def _create(name: str, rows: list[tuple[str, str]]):
        document = Document()
        table = document.add_table(rows=0, cols=2)
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value

        path = tmp_path / name
        document.save(path)
        return path

    return _create


def test_parse_sample_416_extracts_all_fields(docx_factory):
    """Parse file mẫu 416 và kiểm tra toàn bộ field quan trọng được map đúng."""
    path = docx_factory(
        "416.docx",
        """
        Số đăng ký: QN-90599-TS
        Tên chủ tàu: Nguyễn Văn A
        Địa chỉ: Phường Y, Thành phố X
        Tỉnh: Quảng Ninh
        Lmax: 25.5
        Công suất: 500
        Vật liệu: Vỏ Thép
        Hình thức kiểm tra: Hàng năm
        Có giá trị đến: 31/12/2025
        Ngày cấp: 01/01/2024
        Nghề: Lưới kéo
        """,
    )

    vessel = parse_docx(path)

    assert vessel.registration_number == "QN-90599-TS"
    assert vessel.owner_name == "Nguyễn Văn A"
    assert vessel.address == "Phường Y, Thành phố X"
    assert vessel.province_code == "QN"
    assert vessel.lmax == 25.5
    assert vessel.power_kw == 500.0
    assert vessel.material == MaterialEnum.THEP
    assert vessel.inspection_type == InspectionTypeEnum.HANG_NAM
    assert vessel.length_group == LengthGroupEnum.G24_30
    assert vessel.valid_until.isoformat() == "2025-12-31"
    assert vessel.issued_date.isoformat() == "2024-01-01"
    assert vessel.fishing_gear == "Lưới kéo"


def test_parse_sample_417_maps_go_material_and_periodic_inspection(docx_factory):
    """Parse file 417 và kiểm tra vật liệu gỗ, kiểm tra định kỳ, nhóm chiều dài 15-20m."""
    path = docx_factory(
        "417.docx",
        """
        Số đăng ký: TH-12345-TS
        Tên chủ tàu: Trần Văn B
        Địa chỉ: Thanh Hóa
        Tỉnh: Thanh Hóa
        Lmax: 18.0
        Công suất: 300
        Vật liệu: Vỏ Gỗ
        Hình thức kiểm tra: Định kỳ
        """,
    )

    vessel = parse_docx(path)

    assert vessel.registration_number == "TH-12345-TS"
    assert vessel.province_code == "TH"
    assert vessel.material == MaterialEnum.GO
    assert vessel.inspection_type == InspectionTypeEnum.DINH_KY
    assert vessel.length_group == LengthGroupEnum.G15_20


def test_parse_sample_418_maps_frp_and_on_slipway_inspection(docx_factory):
    """Parse file 418 và kiểm tra vật liệu FRP, kiểm tra trên đà, nhóm chiều dài >=30m."""
    path = docx_factory(
        "418.docx",
        """
        Số đăng ký: HP-99999-TS
        Tên chủ tàu: Lê Văn C
        Địa chỉ: Hải Phòng
        Tỉnh: Hải Phòng
        Lmax: 32.0
        Công suất: 800
        Vật liệu: FRP
        Hình thức kiểm tra: Trên đà
        """,
    )

    vessel = parse_docx(path)

    assert vessel.registration_number == "HP-99999-TS"
    assert vessel.province_code == "HP"
    assert vessel.length_group == LengthGroupEnum.G30_PLUS
    assert vessel.material == MaterialEnum.FRP
    assert vessel.inspection_type == InspectionTypeEnum.TREN_DA


def test_missing_required_field_raises_parser_error(docx_factory):
    """File DOCX thiếu số đăng ký phải báo lỗi field bắt buộc thay vì trả dữ liệu rỗng."""
    path = docx_factory(
        "missing-registration.docx",
        """
        Tên chủ tàu: Nguyễn Văn A
        Địa chỉ: Phường Y, Thành phố X
        Lmax: 25.5
        Công suất: 500
        Vật liệu: Gỗ
        Hình thức kiểm tra: Giám sát
        """,
    )

    with pytest.raises(DocxParserError, match="Missing required field"):
        parse_docx(path)


def test_invalid_docx_format_raises_parser_error():
    """File không phải DOCX phải được wrap thành DocxParserError rõ ràng."""
    stream = BytesIO(b"This is not a docx file format")

    with pytest.raises(DocxParserError, match="Invalid DOCX format"):
        parse_docx(stream)


def test_extract_province_code_from_qn_registration_number():
    """Kiểm tra province_code extract đúng từ chuỗi QN-90599-TS."""
    assert extract_province_code("QN-90599-TS") == "QN"


def test_parse_docx_assigns_qn_province_code(docx_factory):
    """Parser phải gán province_code QN từ số đăng ký QN-90599-TS trong dữ liệu vessel."""
    path = docx_factory(
        "province-code.docx",
        """
        Số đăng ký: QN-90599-TS
        Tên chủ tàu: A
        Địa chỉ: B
        Lmax: 12.0
        Công suất: 100
        Vật liệu: Gỗ
        Hình thức kiểm tra: Giám sát
        """,
    )

    vessel = parse_docx(path)

    assert vessel.province_code == "QN"


@pytest.mark.parametrize(
    ("length", "expected"),
    [
        (13.5, LengthGroupEnum.G12_15),
        (22.0, LengthGroupEnum.G20_24),
    ],
)
def test_get_length_group_boundaries(length, expected):
    """Kiểm tra phân loại nhóm chiều dài cho các ngưỡng parser dùng khi tạo VesselData."""
    assert get_length_group(length) == expected


def test_parse_cai_hoan_inspection_type(docx_factory):
    """Parser nhận diện hình thức kiểm tra cải hoán từ nội dung DOCX."""
    path = docx_factory(
        "cai-hoan.docx",
        """
        Số đăng ký: QN-90599-TS
        Tên chủ tàu: A
        Địa chỉ: B
        Lmax: 16.0
        Công suất: 100
        Vật liệu: Gỗ
        Hình thức kiểm tra: Cải hoán
        """,
    )

    vessel = parse_docx(path)

    assert vessel.inspection_type == InspectionTypeEnum.CAI_HOAN


def test_invalid_lmax_format_raises_parser_error(docx_factory):
    """Lmax không có số hợp lệ phải raise DocxParserError để CI bắt được dữ liệu sai."""
    path = docx_factory(
        "invalid-lmax.docx",
        """
        Số đăng ký: QN-90599-TS
        Tên chủ tàu: A
        Địa chỉ: B
        Lmax: ABC
        Công suất: 100
        Vật liệu: Gỗ
        Hình thức kiểm tra: Cải hoán
        """,
    )

    with pytest.raises(DocxParserError, match="Invalid numeric value"):
        parse_docx(path)


def test_invalid_registration_number_format_raises_parser_error(docx_factory):
    """Số đăng ký không có mã tỉnh ở đầu phải raise lỗi format thay vì suy đoán sai."""
    path = docx_factory(
        "invalid-registration-format.docx",
        """
        Số đăng ký: 90599-TS
        Tên chủ tàu: A
        Địa chỉ: B
        Lmax: 16.0
        Công suất: 100
        Vật liệu: Gỗ
        Hình thức kiểm tra: Giám sát
        """,
    )

    with pytest.raises(DocxParserError, match="Invalid registration number format"):
        parse_docx(path)


def test_invalid_material_raises_parser_error(docx_factory):
    """Vật liệu không thuộc danh mục hỗ trợ phải raise lỗi để tránh dữ liệu enum sai."""
    path = docx_factory(
        "invalid-material.docx",
        """
        Số đăng ký: QN-90599-TS
        Tên chủ tàu: A
        Địa chỉ: B
        Lmax: 16.0
        Công suất: 100
        Vật liệu: Xi măng
        Hình thức kiểm tra: Giám sát
        """,
    )

    with pytest.raises(DocxParserError, match="Invalid material"):
        parse_docx(path)


def test_invalid_dates_fallback_to_parser_defaults(docx_factory):
    """Ngày tháng sai format phải fallback về default hiện tại thay vì làm hỏng parse."""
    path = docx_factory(
        "invalid-dates.docx",
        """
        Số đăng ký: HP-111-TS
        Tên chủ tàu: D
        Địa chỉ: E
        Lmax: 15.0
        Công suất: 200
        Vật liệu: Nhựa FRP
        Hình thức kiểm tra: Kiểm tra ABC
        Có giá trị đến: 99/99/9999
        Ngày cấp: InvalidDate
        """,
    )

    vessel = parse_docx(path)

    assert vessel.valid_until.isoformat() == "2024-01-01"
    assert vessel.issued_date.isoformat() == "2020-01-01"
    assert vessel.inspection_type == InspectionTypeEnum.GIAM_SAT


def test_parse_table_based_docx_fields(table_docx_factory):
    """Parser đọc được DOCX có dữ liệu nằm trong bảng key/value, không chỉ paragraph."""
    path = table_docx_factory(
        "table.docx",
        [
            ("Số đăng ký", "BD-77777-TS"),
            ("Tên chủ tàu", "Chủ tàu bảng"),
            ("Địa chỉ", "Bình Định"),
            ("Lmax", "24,5 m"),
            ("Công suất", "450 KW"),
            ("Vật liệu", "Thép"),
            ("Hình thức kiểm tra", "Hàng năm"),
        ],
    )

    vessel = parse_docx(path)

    assert vessel.registration_number == "BD-77777-TS"
    assert vessel.lmax == 24.5
    assert vessel.power_kw == 450.0
    assert vessel.material == MaterialEnum.THEP
