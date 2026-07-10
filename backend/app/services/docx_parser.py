"""DOCX Parser for fishing vessel safety certificates.

Extracts structured vessel data from the standardised DOCX format
used for Vietnamese fishing vessel safety certificates.
"""

import re
from datetime import date
from typing import Optional

from docx import Document
from pydantic import ValidationError

from app.models.vessel import VesselData


class ParseError(Exception):
    """Raised when DOCX parsing fails."""


# Inspection type code mapping
INSPECTION_TYPES = {
    "ĐK": "dinh_ky",
    "HN": "hang_nam",
    "TĐ": "tren_da",
    "GS": "giam_sat",
    "CH": "cai_hoan",
    "BT": "bat_thuong",
    "LĐ": "lan_dau",
    "ĐM": "dong_moi",
}

CLASS_MAP = {
    "Không hạn chế": "khong_han_che",
    "Hạn chế I": "han_che_1",
    "Hạn chế II": "han_che_2",
    "Hạn chế III": "han_che_3",
}

NUMBER_RE = re.compile(r"[-+]?\d+(?:[.,]\d+)?")
DATE_RE = re.compile(
    r"ngày\s*(\d{1,2})\s*tháng\s*(\d{1,2})\s*năm\s*(\d{4})", re.IGNORECASE
)


def normalize_text(text: str) -> str:
    """Collapse whitespace and strip."""
    return re.sub(r"\s+", " ", text or "").strip()


def strip_english_suffix(text: str) -> str:
    """Remove trailing English labels commonly found in bilingual cells."""
    text = re.sub(r"^\s*\(.*?\)\s*:?\s*", "", text)
    text = re.sub(r"[.…]+\s*$", "", text)
    text = re.sub(r"\s*\([^)]*\)\s*$", "", text)
    text = re.sub(
        r"\s+\b(owner|representative|captain|co-owner|coowner|manager|address|nationality|used\s+for|used|for|fishing|trade|service|gear)\b\s*$",
        "",
        text,
        flags=re.IGNORECASE
    )
    return text.strip(". \t")


def extract_number(text: str) -> Optional[float]:
    """Extract the first number from text, handling Vietnamese comma decimals."""
    m = NUMBER_RE.search(normalize_text(text))
    if not m:
        return None
    return float(m.group(0).replace(",", "."))


def extract_date_vn(text: str) -> Optional[date]:
    """Extract Vietnamese date ngày X tháng Y năm Z → date object."""
    m = DATE_RE.search(normalize_text(text))
    if not m:
        return None
    try:
        return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
    except ValueError:
        return None


def extract_owner_name(text: str) -> str:
    """Extract the owner's full name from a certificate paragraph or full document text."""
    text = normalize_text(text)
    owner_pattern = re.search(
        r"(?:Chủ tàu|Họ và tên|[Oo]wner)(?:\s*\([^)]+\))?\s*[:\-–—]?\s*(.+?)"
        r"(?=\s*(?:;|Quốc tịch\b|Nationality\b|Địa chỉ\b|Address\b|\(\s*[^)]+\s*\)\s*:|$))",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if owner_pattern:
        owner_value = normalize_text(owner_pattern.group(1))
        owner_value = strip_english_suffix(owner_value)
        return owner_value.strip(" :;\t\n")

    fallback_pattern = re.search(
        r"(?:Chủ tàu|Họ và tên)[^:]*:[^:]*:?[\s\S]*?(.+?)"
        r"(?=\s*(?:;|Quốc tịch\b|Nationality\b|Địa chỉ\b|Address\b|\(\s*[^)]+\s*\)\s*:|$))",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if fallback_pattern:
        owner_value = normalize_text(fallback_pattern.group(1))
        owner_value = strip_english_suffix(owner_value)
        return owner_value.strip(" :;\t\n")

    return ""


def extract_build_info(text: str) -> tuple[Optional[str], Optional[str]]:
    """Extract build year and build place from text."""
    # Matches "Năm và nơi đóng"
    m = re.search(r"Năm và nơi đóng[^:]*:\s*([^|;\n]+(?:[;,][^|;\n]+)*)", text, re.IGNORECASE)
    if not m:
        return None, None
    raw_info = m.group(1).strip()
    raw_info = strip_english_suffix(raw_info)
    
    parts = [p.strip() for p in re.split(r'[;,]', raw_info) if p.strip()]
    year = None
    place = None
    
    # Check if any part is a 4-digit number
    for i, part in enumerate(parts):
        if re.match(r"^\d{4}$", part):
            year = part
            other_parts = parts[:i] + parts[i+1:]
            place = "; ".join(other_parts) if other_parts else None
            break
            
    if not year and parts:
        if len(parts) >= 2:
            year = parts[0]
            place = "; ".join(parts[1:])
        else:
            if re.search(r"\d{4}", parts[0]):
                year = parts[0]
            else:
                place = parts[0]
                
    return year, place


def extract_gross_tonnage(text: str) -> Optional[float]:
    """Extract gross tonnage from text."""
    m = re.search(r"Tổng dung tích[^:\d]*:\s*(\d+(?:[.,]\d+)?)", text, re.IGNORECASE)
    if m:
        val_str = m.group(1).replace(",", ".")
        try:
            return float(val_str)
        except ValueError:
            return None
    return None


def extract_crew_limit(text: str) -> Optional[int]:
    """Extract crew limit from text."""
    m = re.search(r"Số thuyền viên[^:\d]*:\s*(\d+)", text, re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except ValueError:
            return None
    return None


def extract_bmax(text: str) -> Optional[float]:
    """Extract Bmax from text."""
    m = re.search(r"(?:Chiều rộng[^,]*,\s*Bmax|Bmax)[^:\d]*:\s*(\d+(?:[.,]\d+)?)", text, re.IGNORECASE)
    if m:
        val_str = m.group(1).replace(",", ".")
        try:
            return float(val_str)
        except ValueError:
            return None
    return None


def extract_depth(text: str) -> Optional[float]:
    """Extract depth overall from text."""
    m = re.search(r"(?:Chiều cao mạn[^,]*,\s*D|Chiều cao mạn)[^:\d]*:\s*(\d+(?:[.,]\d+)?)", text, re.IGNORECASE)
    if m:
        val_str = m.group(1).replace(",", ".")
        try:
            return float(val_str)
        except ValueError:
            return None
    return None


def extract_allowed_area(text: str) -> Optional[str]:
    """Extract allowed operating area from text."""
    m = re.search(r"Được phép hoạt động tại[^:]*:\s*(.+?)(?=\s*(?:Giấy chứng nhận|This certificate\b|ngày\b|\||$))", text, re.IGNORECASE)
    if m:
        val = m.group(1).strip()
        val = strip_english_suffix(val)
        return val.strip(" :;\t\n")
    return ""


def safe_cell_text(table, row_idx: int, col_idx: int) -> str:
    """Get cell text safely, returning '' if out of bounds."""
    try:
        if row_idx < len(table.rows):
            cells = table.rows[row_idx].cells
            if col_idx < len(cells):
                return normalize_text(cells[col_idx].text)
    except (IndexError, AttributeError):
        pass
    return ""


def _build_full_text(doc) -> str:
    """Build concatenated full document text for fallback regex searches."""
    blocks: list[str] = []
    for para in doc.paragraphs:
        t = normalize_text(para.text)
        if t:
            blocks.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = normalize_text(cell.text)
                if t:
                    blocks.append(t)
    return " | ".join(blocks)


def format_address_short(dia_chi: str, ma_tinh: str) -> str:
    """Create short address format like 'NA, Diễn Châu'"""
    if not dia_chi:
        return ma_tinh
    parts = dia_chi.split(",")
    first_part = parts[0].strip()
    # Remove common prefixes
    for prefix in ["xã", "Xã", "huyện", "Huyện", "thị xã", "Thị xã", "thành phố", "Thành phố", "tp", "TP"]:
        if first_part.startswith(prefix):
            first_part = first_part[len(prefix):].strip()
            break
    if first_part:
        return f"{ma_tinh}, {first_part}"
    return ma_tinh


def parse_vessel_docx(file_path: str) -> VesselData:
    """Parse a fishing vessel safety certificate DOCX file."""
    try:
        doc = Document(file_path)
    except Exception as exc:
        raise ParseError(f"Không thể đọc file DOCX: {exc}")

    tables = doc.tables
    paragraphs = doc.paragraphs
    full_text = _build_full_text(doc)

    # 1. Inspection type
    hinh_thuc_kiem_tra = "khong_xac_dinh"
    cert_text = safe_cell_text(tables[0], 0, 0) if tables else ""
    
    cert_match = re.search(
        r"Số[^:\d]*/?\s*[\d.]+\s*/\s*([a-zA-ZĐđ]+)", cert_text, re.IGNORECASE
    )
    if not cert_match:
        cert_match = re.search(
            r"Số[^:\d]*/?\s*[\d.]+\s*/\s*([a-zA-ZĐđ]+)", full_text, re.IGNORECASE
        )
    if not cert_match:
        cert_match = re.search(
            r"[\d.]+\s*/\s*([a-zA-ZĐđ]{2})", cert_text, re.IGNORECASE
        )
    if not cert_match:
        cert_match = re.search(
            r"[\d.]+\s*/\s*([a-zA-ZĐđ]{2})", full_text, re.IGNORECASE
        )
        
    if cert_match:
        ma_hk = cert_match.group(1).upper()
        hinh_thuc_kiem_tra = INSPECTION_TYPES.get(ma_hk, "khong_xac_dinh")

    # 2. Registration number (extracts province + number)
    so_dang_ky = None
    ma_tinh = "UNK"
    reg_pattern = re.compile(r"([a-zA-ZĐđ]{2,4})\s*-\s*([^-\s]+)\s*-\s*[Tt][Ss]")
    
    reg_text = ""
    if len(tables) > 1:
        reg_text = safe_cell_text(tables[1], 0, 1)
    
    reg_match = reg_pattern.search(reg_text)
    if not reg_match:
        reg_match = reg_pattern.search(full_text)
        
    if reg_match:
        ma_tinh = reg_match.group(1).strip()
        num_part = reg_match.group(2).strip()
        if re.search(r"\d", num_part):
            so_dang_ky = f"{ma_tinh}-{num_part}-TS".upper()
        else:
            so_dang_ky = None # Missing number (e.g. TH-....-TS)

    # 3. Owner name
    ho_ten = ""
    if len(paragraphs) > 2:
        ho_ten = extract_owner_name(normalize_text(paragraphs[2].text))
    if not ho_ten:
        ho_ten = extract_owner_name(full_text)

    # 4. Address
    dia_chi = ""
    if len(paragraphs) > 3:
        p3 = normalize_text(paragraphs[3].text)
        m = re.search(r"Địa chỉ[^:]*:\s*(?:\(Address\)\s*:?\s*)?(.+)", p3, re.IGNORECASE)
        if m:
            dia_chi = m.group(1).strip()
    if not dia_chi:
        m = re.search(r"Địa chỉ[^:]*:\s*([^;\n|]+)", full_text, re.IGNORECASE)
        if m:
            raw = m.group(1).strip()
            if raw.lower().startswith("(address)"):
                raw = raw[9:].strip(": \t")
            dia_chi = raw
            
    address_short = format_address_short(dia_chi, ma_tinh.upper().replace('Đ', 'Đ'))

    # 5. Technical specs
    nghe = ""
    if len(tables) > 2:
        t2r0c0 = safe_cell_text(tables[2], 0, 0)
        m = re.search(r"Công dụng[^:]*:\s*([^(\n]+)", t2r0c0, re.IGNORECASE)
        if m:
            nghe = strip_english_suffix(m.group(1).strip())
    if not nghe:
        m = re.search(
            r"(?:Công dụng|Nghề)[^:]*:\s*(.+?)(?=\s*(?:;|\n|\||\b(?:Use|Trade)\b|\(|$))",
            full_text,
            re.IGNORECASE
        )
        if m:
            nghe = m.group(1).strip()

    vat_lieu = "Không xác định"
    if len(tables) > 2 and tables[2].rows:
        for cell in reversed(list(tables[2].rows[0].cells)):
            ct = normalize_text(cell.text)
            m = re.search(r"Vật liệu[^:]*:\s*([^(\n]+)", ct, re.IGNORECASE)
            if m:
                mat = strip_english_suffix(m.group(1).strip()).lower()
                if "gỗ" in mat or "go" in mat: vat_lieu = "Gỗ"
                elif "thép" in mat or "thep" in mat: vat_lieu = "Thép"
                elif "frp" in mat or "composite" in mat: vat_lieu = "FRP"
                break
    if vat_lieu == "Không xác định":
        m = re.search(r"Vật liệu[^:]*:\s*([^(\n|]+)", full_text, re.IGNORECASE)
        if m:
            mat = m.group(1).strip().lower()
            if "gỗ" in mat or "go" in mat: vat_lieu = "Gỗ"
            elif "thép" in mat or "thep" in mat: vat_lieu = "Thép"
            elif "frp" in mat or "composite" in mat: vat_lieu = "FRP"

    lmax = 0.0
    if len(tables) > 2:
        t2r2c0 = safe_cell_text(tables[2], 2, 0)
        if "Lmax" in t2r2c0 or "lmax" in t2r2c0.lower():
            after_label = t2r2c0.split("Lmax")[-1]
            val = extract_number(re.sub(r"\(.*?\)", "", after_label))
            if val and val > 0: lmax = val
    if lmax <= 0:
        m = re.search(r"(?:Chiều dài[^,]*,\s*Lmax|Lmax)[^:]*:\s*([^(\n|]+)", full_text, re.IGNORECASE)
        if m:
            val = extract_number(m.group(1))
            if val and val > 0: lmax = val

    may_chinh = 0.0
    if len(tables) > 2:
        t2r4c0 = safe_cell_text(tables[2], 4, 0)
        m = re.search(r"Ne\s*\(?KW\)?[^:]*:\s*([^(\n]+)", t2r4c0, re.IGNORECASE)
        if m:
            val = extract_number(m.group(1))
            if val: may_chinh = val
    if may_chinh <= 0:
        m = re.search(r"(?:công suất máy chính|Công suất|Ne)[^:]*:\s*([^(\n|]+)", full_text, re.IGNORECASE)
        if m:
            val = extract_number(m.group(1))
            if val: may_chinh = val

    # 5.5. Additional Technical specs
    build_year, build_place = extract_build_info(full_text)
    gross_tonnage = extract_gross_tonnage(full_text)
    crew_limit = extract_crew_limit(full_text)
    bmax = extract_bmax(full_text)
    depth = extract_depth(full_text)
    allowed_area = extract_allowed_area(full_text)

    # Main Engine details table
    engine_model = ""
    engine_serial = ""
    engine_build_info = ""
    
    engine_table = None
    for table in tables:
        if table.rows:
            header_text = "".join(normalize_text(cell.text) for cell in table.rows[0].cells)
            if "Ký hiệu máy" in header_text or "Type of machine" in header_text or "Số máy" in header_text:
                engine_table = table
                break
                
    if engine_table and len(engine_table.rows) > 1:
        r1 = engine_table.rows[1]
        if len(r1.cells) >= 5:
            engine_model = strip_english_suffix(normalize_text(r1.cells[1].text))
            engine_serial = strip_english_suffix(normalize_text(r1.cells[2].text))
            engine_build_info = strip_english_suffix(normalize_text(r1.cells[4].text))
            
            if engine_model == "-": engine_model = ""
            if engine_serial == "-": engine_serial = ""
            if engine_build_info == "-": engine_build_info = ""

    # 6. Operation class
    cap_tau = "khong_xac_dinh"
    if len(tables) > 4:
        t4_rows = [[normalize_text(c.text) for c in row.cells] for row in tables[4].rows]
        if len(t4_rows) >= 2:
            header, values = t4_rows[0], t4_rows[1]
            for cap_name, cap_code in CLASS_MAP.items():
                for idx, cell_val in enumerate(header):
                    if cap_name in cell_val and idx < len(values):
                        if values[idx].strip().upper() in ("X", "x"):
                            cap_tau = cap_code
                            break
                if cap_tau != "khong_xac_dinh":
                    break

    # 6.5. Technical status
    trang_thai_kt = ""
    if len(paragraphs) > 10:
        p10 = normalize_text(paragraphs[10].text)
        if ":" in p10:
            trang_thai_kt = p10.split(":")[-1].strip()
            trang_thai_kt = strip_english_suffix(trang_thai_kt)
    if not trang_thai_kt:
        m = re.search(r"Trạng\s*thái\s*k[ỹĩ]\s*thuật[^:]*:\s*([^;|]+)", full_text, re.IGNORECASE)
        if m:
            raw = m.group(1).strip()
            if raw.lower().startswith("(technical status)"):
                raw = raw[18:].strip(": \t")
            trang_thai_kt = strip_english_suffix(raw)

    # 7. Dates
    # Inspection date
    ngay_kt = None
    m_kt = re.search(r"biên bản kiểm tra[^ngày]*ngày\s*\d{1,2}\s*tháng\s*\d{1,2}\s*năm\s*\d{4}", full_text, re.IGNORECASE)
    if m_kt:
        ngay_kt = extract_date_vn(m_kt.group(0))

    # Valid until
    han_dk = None
    m_han = re.search(r"(?:có hiệu lực đến|giá trị đến hết ngày|Hạn đăng kiểm)[^\n|]*", full_text, re.IGNORECASE)
    if m_han:
        han_dk = extract_date_vn(m_han.group(0))

    # Issued date
    ngay_cap = None
    if len(tables) > 5:
        ngay_cap = extract_date_vn(safe_cell_text(tables[5], 0, 1))
    if not ngay_cap and tables:
        for row in tables[-1].rows:
            for cell in row.cells:
                d = extract_date_vn(normalize_text(cell.text))
                if d:
                    ngay_cap = d
                    break
            if ngay_cap: break
            
    if not ngay_kt:
        ngay_kt = ngay_cap or date.today()

    try:
        import os
        return VesselData(
            registration_no=so_dang_ky,
            province_code=ma_tinh,
            owner_name=ho_ten,
            address=dia_chi,
            address_short=address_short,
            lmax=lmax,
            power_kw=may_chinh,
            material=vat_lieu,
            inspection_type=hinh_thuc_kiem_tra,
            vessel_class=cap_tau,
            technical_status=trang_thai_kt,
            inspection_date=ngay_kt,
            valid_until=han_dk,
            issued_date=ngay_cap,
            fishing_gear=nghe,
            source_filename=os.path.basename(file_path),
            build_year=build_year,
            build_place=build_place,
            gross_tonnage=gross_tonnage,
            crew_limit=crew_limit,
            bmax=bmax,
            depth=depth,
            engine_model=engine_model,
            engine_serial=engine_serial,
            engine_build_info=engine_build_info,
            allowed_area=allowed_area,
        )
    except ValidationError as exc:
        raise ParseError(f"Lỗi xác thực dữ liệu: {exc}")
